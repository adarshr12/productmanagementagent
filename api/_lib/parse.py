"""Extract plain text from an uploaded file (PDF / DOCX / TXT).

Uses the open-source libraries pypdf and python-docx. Everything happens in
memory — we never write the raw file to disk or storage.
"""
from __future__ import annotations

import io


def extract_text(source_type: str, data: bytes) -> str:
    """Return the plain text of a file given its raw bytes and type."""
    source_type = (source_type or "").lower()

    if source_type == "pdf":
        return _from_pdf(data)
    if source_type == "docx":
        return _from_docx(data)
    if source_type == "txt":
        return data.decode("utf-8", errors="replace")

    raise ValueError(f"Unsupported source_type: {source_type!r}")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _from_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include table cell text too — resource lists are often in tables.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)
